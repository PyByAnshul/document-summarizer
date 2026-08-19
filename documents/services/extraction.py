import logging
from pathlib import Path

import pymupdf as fitz
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentExtractionError(Exception):
    pass


def extract_from_pdf(file_path):
    file_name = Path(file_path).name
    logger.info("PDF text extraction started: file=%s", file_name)
    try:
        with fitz.open(file_path) as pdf:
            text = "\n".join(page.get_text() for page in pdf)
            logger.info(
                "PDF text extraction completed: file=%s pages=%d characters=%d",
                file_name,
                len(pdf),
                len(text),
            )
            return text
    except Exception as exc:
        logger.exception("PDF text extraction failed: file=%s", file_name)
        raise DocumentExtractionError("Failed to extract text from PDF.") from exc


def extract_from_docx(file_path):
    file_name = Path(file_path).name
    logger.info("DOCX text extraction started: file=%s", file_name)
    try:
        document = DocxDocument(file_path)
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        logger.info(
            "DOCX text extraction completed: file=%s paragraphs=%d characters=%d",
            file_name,
            len(document.paragraphs),
            len(text),
        )
        return text
    except Exception as exc:
        logger.exception("DOCX text extraction failed: file=%s", file_name)
        raise DocumentExtractionError("Failed to extract text from DOCX.") from exc


def extract_from_txt(file_path):
    file_name = Path(file_path).name
    logger.info("TXT text extraction started: file=%s", file_name)
    try:
        with open(file_path, encoding="utf-8") as file:
            text = file.read()
        logger.info(
            "TXT text extraction completed: file=%s characters=%d",
            file_name,
            len(text),
        )
        return text
    except Exception as exc:
        logger.exception("TXT text extraction failed: file=%s", file_name)
        raise DocumentExtractionError("Failed to extract text from TXT.") from exc


def extract_text(file_path):
    extension = Path(file_path).suffix.lower()
    logger.info("Document extraction requested: file=%s type=%s", Path(file_path).name, extension)

    extractors = {
        ".pdf": extract_from_pdf,
        ".docx": extract_from_docx,
        ".txt": extract_from_txt,
    }

    if extension not in extractors:
        raise ValueError(
            "Unsupported file type. Only PDF, DOCX and TXT files are supported."
        )

    text = extractors[extension](file_path)

    if not text.strip():
        logger.warning("Document extraction produced no text: file=%s", Path(file_path).name)
        raise ValueError("Document text is empty.")

    logger.info(
        "Document extraction finished: file=%s words=%d",
        Path(file_path).name,
        len(text.split()),
    )
    return text
